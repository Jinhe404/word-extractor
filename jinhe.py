import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document


def extract_content_between_keywords(file_path, start_keyword, end_keyword):
    doc = Document(file_path)
    found_start = False
    content_between_keywords = []

    def process_paragraphs(paragraphs):
        nonlocal found_start
        for para in paragraphs:
            if start_keyword in para.text:
                found_start = True
                continue
            if found_start:
                content_between_keywords.append(para.text)
                if end_keyword in para.text:
                    return True
        return False

    if not process_paragraphs(doc.paragraphs):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if process_paragraphs(cell.paragraphs):
                        return '\n'.join(content_between_keywords)

    return '\n'.join(content_between_keywords)


def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        file_var.set(file_path)


def extract_and_save():
    file_path = file_var.get()
    start_keyword = start_keyword_var.get()
    end_keyword = end_keyword_var.get()

    if not file_path or not start_keyword or not end_keyword:
        messagebox.showerror("错误", "请填写所有字段")
        return

    content = extract_content_between_keywords(file_path, start_keyword, end_keyword)
    save_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
    if save_path:
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(content)
        messagebox.showinfo("成功", "内容已保存到文件")


root = tk.Tk()
root.title("Word内容提取器")

file_var = tk.StringVar()
start_keyword_var = tk.StringVar()
end_keyword_var = tk.StringVar()

tk.Label(root, text="Word文件路径:").grid(row=0, column=0, padx=10, pady=5)
tk.Entry(root, textvariable=file_var, width=50).grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="打开", command=open_file).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="开始关键词:").grid(row=1, column=0, padx=10, pady=5)
tk.Entry(root, textvariable=start_keyword_var).grid(row=1, column=1, padx=10, pady=5)

tk.Label(root, text="结束关键词:").grid(row=2, column=0, padx=10, pady=5)
tk.Entry(root, textvariable=end_keyword_var).grid(row=2, column=1, padx=10, pady=5)

tk.Button(root, text="开始截取", command=extract_and_save).grid(row=3, column=0, columnspan=3, pady=20)

root.mainloop()
